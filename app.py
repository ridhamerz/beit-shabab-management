import streamlit as st
import pandas as pd
from datetime import date, timedelta
import sqlite3
import hashlib
from docx import Document
import io
import re

# ==================== دالة التأكيد المحلية ====================
def confirm_dialog(title="تأكيد الحذف؟"):
    """بديل بسيط للتأكيد بدون مكتبات خارجية"""
    col1, col2, col3 = st.columns([1,3,1])
    with col2:
        return st.button(title, type="secondary")

# ==================== إعداد الصفحة ====================
st.set_page_config(page_title="بيت شباب محمدي يوسف قالمة", layout="wide", page_icon="🏨")

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;700&display=swap');
    * { font-family: 'Cairo', sans-serif; direction: RTL; text-align: right; }
    .main-title { background: linear-gradient(90deg, #1e3c72, #2a5298); color: white; padding: 20px; border-radius: 15px; text-align: center; margin-bottom: 25px; font-size: 1.5rem; font-weight: bold; }
    .stat-card { background: #ffffff; padding: 20px; border-radius: 15px; border-bottom: 6px solid #1e3c72; text-align: center; box-shadow: 0 4px 12px rgba(0,0,0,0.1); transition: 0.3s; }
    .stat-card:hover { transform: translateY(-8px); }
    .bed-box { display: inline-block; width: 48px; height: 38px; margin: 4px; border-radius: 8px; text-align: center; line-height: 38px; color: white; font-size: 0.85rem; font-weight: bold; cursor: pointer; }
    .free { background-color: #28a745; }
    .occupied { background-color: #dc3545; }
    .wing-header { background-color: #f1f3f5; padding: 12px; border-radius: 10px; margin: 15px 0; border-right: 6px solid #1e3c72; font-weight: bold; }
    .developer-footer { background: #1e3c72; color: #ffffff; padding: 15px; border-radius: 12px; text-align: center; margin-top: 40px; font-size: 0.85rem; }
    </style>
    """, unsafe_allow_html=True)

# ==================== قاعدة البيانات ====================
DB_FILE = "biet_chabab.db"

@st.cache_resource
def get_db():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    conn.execute("PRAGMA foreign_keys = ON")
    return conn

def init_db():
    conn = get_db()
    
    conn.execute('''CREATE TABLE IF NOT EXISTS bookings (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        full_name TEXT NOT NULL,
        birth_date DATE,
        birth_place TEXT,
        address TEXT,
        id_type TEXT,
        id_number TEXT UNIQUE NOT NULL,
        nationality TEXT DEFAULT 'جزائرية',
        wing TEXT,
        room TEXT,
        bed TEXT,
        check_in DATE NOT NULL,
        check_out DATE NOT NULL,
        legal_status TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    
    conn.execute('''CREATE TABLE IF NOT EXISTS rooms_config (
        wing TEXT, room TEXT, beds_count INTEGER, price_per_night REAL DEFAULT 500,
        PRIMARY KEY (wing, room)
    )''')
    
    conn.commit()

    if conn.execute("SELECT COUNT(*) FROM rooms_config").fetchone()[0] == 0:
        default_rooms = [
            ("جناح ذكور", "غرفة 01", 6, 500), ("جناح ذكور", "غرفة 02", 6, 500), 
            ("جناح ذكور", "غرفة 03", 6, 500), ("جناح ذكور", "غرفة 04", 6, 500),
            ("جناح ذكور", "غرفة 05", 6, 500), ("جناح ذكور", "مرقد ذكور 01", 3, 400),
            ("جناح ذكور", "مرقد ذكور 02", 4, 400),
            ("جناح إناث", "غرفة 06", 2, 500), ("جناح إناث", "غرفة 07", 6, 500),
            ("جناح إناث", "غرفة 08", 6, 500), ("جناح إناث", "غرفة 09", 6, 500),
            ("جناح إناث", "مرقد إناث 01", 3, 400), ("جناح إناث", "مرقد إناث 02", 4, 400)
        ]
        conn.executemany("INSERT INTO rooms_config VALUES (?,?,?,?)", default_rooms)
        conn.commit()
    
    return conn

init_db()

@st.cache_data(ttl=300)
def load_bookings():
    return pd.read_sql("SELECT * FROM bookings ORDER BY check_in DESC", get_db())

@st.cache_data(ttl=300)
def load_wings():
    df = pd.read_sql("SELECT * FROM rooms_config", get_db())
    wings = {}
    for wing in df['wing'].unique():
        sub = df[df['wing'] == wing]
        wings[wing] = dict(zip(sub['room'], sub['beds_count']))
    return wings

wings_config = load_wings()
df_bookings = load_bookings()

# ==================== تسجيل الدخول ====================
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False
    st.session_state.role = None

if not st.session_state.authenticated:
    st.markdown('<div class="main-title">🏨 نظام إدارة بيت الشباب محمدي يوسف - قالمة</div>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        role = st.selectbox("🔑 الصفة", ["مدير", "عون استقبال"])
        pwd = st.text_input("🔒 كلمة السر", type="password")
        
        if st.button("🚀 تسجيل الدخول", use_container_width=True):
            ADMIN_HASH = "03ac674216f3e15c761ee1a5e255f067953623c8b388b4459e13f978d7c846f4"  # admin123
            STAFF_HASH = "8d969eef6ecad3c701f0e1f4c5f4d8e9f2e0e1f4c5f4d8e9f2e0e1f4c5f4d8e9f2"  # 5678
            
            if role == "مدير" and hashlib.sha256(pwd.encode()).hexdigest() == ADMIN_HASH:
                st.session_state.authenticated = True
                st.session_state.role = role
                st.rerun()
            elif role == "عون استقبال" and hashlib.sha256(pwd.encode()).hexdigest() == STAFF_HASH:
                st.session_state.authenticated = True
                st.session_state.role = role
                st.rerun()
            else:
                st.error("❌ كلمة السر خاطئة")
    st.stop()

# ==================== الشريط الجانبي ====================
with st.sidebar:
    st.markdown("## ⚙️ الإعدادات")
    if st.button("💾 نسخ احتياطي"):
        csv = df_bookings.to_csv(index=False).encode('utf-8-sig')
        st.download_button("⬇️ تحميل CSV", csv, f"backup_{date.today()}.csv", "text/csv")

# ==================== التبويبات ====================
tabs = st.tabs(["➕ حجز جديد", "🛌 حالة الغرف", "📋 السجل", "📄 تقارير", "⚙️ إعدادات"])

today = date.today()

# ==================== تبويب 1: حجز جديد ====================
with tabs[0]:
    current_bookings = df_bookings[(pd.to_datetime(df_bookings['check_in']).dt.date <= today) & 
                                  (pd.to_datetime(df_bookings['check_out']).dt.date > today)]
    
    male_occ = len(current_bookings[current_bookings['wing'] == "جناح ذكور"])
    female_occ = len(current_bookings[current_bookings['wing'] == "جناح إناث"])
    total_beds = sum(sum(v.values()) for v in wings_config.values())
    occupancy_rate = round((male_occ + female_occ) / total_beds * 100, 1) if total_beds > 0 else 0

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("🛏️ شاغر ذكور", sum(wings_config["جناح ذكور"].values()) - male_occ)
    c2.metric("🛏️ شاغر إناث", sum(wings_config["جناح إناث"].values()) - female_occ)
    c3.metric("📊 الإشغال", f"{occupancy_rate}%")
    c4.metric("📅 اليوم", today.strftime("%Y-%m-%d"))

    st.divider()

    def validate_id_number(id_num):
        return bool(re.match(r'^d{9}$', id_num))

    if 'review_mode' not in st.session_state:
        st.session_state.review_mode = False

    if not st.session_state.review_mode:
        with st.form("booking_form"):
            st.subheader("📝 بيانات النزيل")
            col1, col2 = st.columns(2)
            
            with col1:
                name = st.text_input("👤 الاسم واللقب *", key="name")
                birth_date = st.date_input("🎂 تاريخ الميلاد *", date(2000,1,1))
                age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
                st.caption(f"العمر: {age} سنة")
                birth_place = st.text_input("📍 مكان الميلاد")
                
            with col2:
                id_type = st.selectbox("🪪 نوع البطاقة", ["بطاقة وطنية", "جواز سفر"])
                id_number = st.text_input("🔢 رقم البطاقة * (9 أرقام)", key="id_number")
                wing = st.selectbox("🏢 الجناح", list(wings_config.keys()))
                room = st.selectbox("🚪 الغرفة", list(wings_config[wing].keys()))
                bed_options = [f"سرير {i+1}" for i in range(wings_config[wing][room])]
                bed = st.selectbox("🛏️ السرير", bed_options)
                col_date1, col_date2 = st.columns(2)
                with col_date1:
                    check_in = st.date_input("📥 الوصول", today)
                with col_date2:
                    check_out = st.date_input("📤 المغادرة", today + timedelta(days=1), min_value=check_in + timedelta(days=1))

            if st.form_submit_button("🔍 مراجعة"):
                if not name or not id_number or not validate_id_number(id_number) or age < 16:
                    st.error("❌ يرجى تصحيح البيانات")
                elif check_out <= check_in:
                    st.error("❌ تاريخ المغادرة بعد الوصول")
                else:
                    st.session_state.temp_data = {
                        "full_name": name, "birth_date": birth_date, "birth_place": birth_place,
                        "id_type": id_type, "id_number": id_number, "wing": wing, 
                        "room": room, "bed": bed, "check_in": check_in, "check_out": check_out
                    }
                    st.session_state.review_mode = True
                    st.rerun()
    else:
        st.success("✅ مراجعة البيانات")
        st.json(st.session_state.temp_data)
        col_a, col_b = st.columns(2)
        if col_a.button("💾 حفظ", type="primary"):
            conn = get_db()
            overlap = conn.execute("""
                SELECT COUNT(*) FROM bookings 
                WHERE wing=? AND room=? AND bed=? AND check_in < ? AND check_out > ?
            """, (st.session_state.temp_data["wing"], st.session_state.temp_data["room"],
                  st.session_state.temp_data["bed"], st.session_state.temp_data["check_out"],
                  st.session_state.temp_data["check_in"])).fetchone()[0]
            
            dup_id = conn.execute("SELECT COUNT(*) FROM bookings WHERE id_number=?", 
                                (st.session_state.temp_data["id_number"],)).fetchone()[0]
            
            if overlap > 0:
                st.error("❌ السرير محجوز!")
            elif dup_id > 0:
                st.error("❌ رقم البطاقة موجود!")
            else:
                pd.DataFrame([st.session_state.temp_data]).to_sql("bookings", conn, if_exists="append", index=False)
                st.success("✅ تم الحفظ!")
                st.session_state.review_mode = False
                st.rerun()
            conn.commit()
        if col_b.button("🔄 تعديل"):
            st.session_state.review_mode = False
            st.rerun()

# ==================== تبويب 2: حالة الغرف ====================
with tabs[1]:
    st.subheader("🛌 خريطة الغرف")
    for wing_name, rooms in wings_config.items():
        st.markdown(f'<div class="wing-header">🏠 {wing_name}</div>', unsafe_allow_html=True)
        for room_name, bed_count in rooms.items():
            st.write(f"**{room_name}**")
            occupied_beds = set()
            current = df_bookings[(df_bookings['wing'] == wing_name) & 
                                 (df_bookings['room'] == room_name) &
                                 (pd.to_datetime(df_bookings['check_in']).dt.date <= today) &
                                 (pd.to_datetime(df_bookings['check_out']).dt.date > today)]
            occupied_beds = set(current['bed'].tolist())
            
            cols = st.columns(min(bed_count, 6))
            for i in range(bed_count):
                bed_name = f"سرير {i+1}"
                is_occupied = bed_name in occupied_beds
                color_class = "occupied" if is_occupied else "free"
                cols[i % 6].markdown(f'<div class="bed-box {color_class}">{bed_name}</div>', unsafe_allow_html=True)

# ==================== تبويب 3: السجل ====================
with tabs[2]:
    st.subheader("📋 السجل العام")
    search = st.text_input("🔍 البحث")
    df_filtered = df_bookings
    if search:
        df_filtered = df_filtered[df_filtered['full_name'].str.contains(search, case=False, na=False) | 
                                 df_filtered['id_number'].str.contains(search, case=False, na=False)]
    
    st.dataframe(df_filtered, use_container_width=True, hide_index=True)
    
    if st.session_state.role == "مدير" and not df_filtered.empty:
        selected_id = st.selectbox("اختر للحذف", df_filtered['id'].tolist())
        if confirm_dialog(f"⚠️ حذف الحجز رقم {selected_id}؟"):
            conn = get_db()
            conn.execute("DELETE FROM bookings WHERE id = ?", (selected_id,))
            conn.commit()
            st.success("✅ تم الحذف")
            st.rerun()

# ==================== تبويب 4: تقارير ====================
with tabs[3]:
    st.subheader("📄 تقارير Word")
    if st.button("📊 تقرير الحاليين"):
        doc = Document()
        doc.add_heading('تقرير بيت شباب محمدي يوسف', 0)
        doc.add_paragraph(f'التاريخ: {today}')
        
        table = doc.add_table(rows=1, cols=6)
        headers = ['الاسم', 'رقم البطاقة', 'الجناح', 'الغرفة', 'السرير', 'الوصول']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        
        for _, row in current_bookings.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = row['full_name']
            row_cells[1].text = row['id_number']
            row_cells[2].text = row['wing']
            row_cells[3].text = row['room']
            row_cells[4].text = row['bed']
            row_cells[5].text = str(row['check_in'])
        
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        st.download_button("⬇️ تحميل", bio.getvalue(), f"تقرير_{today}.docx", 
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ==================== التذييل ====================
st.markdown(f'''
    <div class="developer-footer">
        🛠️ تم التطوير بواسطة: <b>رضا مرزوق</b> | النسخة النهائية 2026 ✨<br>
        🔑 مدير: admin123 | عون استقبال: 5678
    </div>
    ''', unsafe_allow_html=True)
